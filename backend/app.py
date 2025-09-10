from flask import Flask, request, send_file, jsonify
import pikepdf # Use pikepdf for PDF operations
from flask_cors import CORS
import io
import logging
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz  # PyMuPDF for better text extraction
from PIL import Image, ImageOps, ImageEnhance  # Add Pillow imports for image processing

# Initialize Flask app
app = Flask(__name__)
CORS(app)  # Enable CORS for cross-origin requests

# Configure logging
logging.basicConfig(level=logging.INFO) # Set to INFO for production, DEBUG for development



# Root route and health endpoint
@app.route('/')
def home():
    return "PDF Manipulator Backend is running!"

@app.route('/health')
def health():
    return jsonify({"status": "ok"})

# Unlock PDF endpoint
@app.route('/unlock-pdf', methods=['POST'])
def unlock_pdf():
    # Check for file and password in request
    if 'file' not in request.files:
        logging.error("Unlock PDF: No file part in the request.")
        return jsonify({"error": "No file part in the request."}), 400
    if 'password' not in request.form:
        logging.error("Unlock PDF: Password not provided.")
        return jsonify({"error": "Password not provided."}), 400

    file = request.files['file']
    password = request.form.get('password')

    # Validate file presence and type
    if file.filename == '':
        logging.error("Unlock PDF: No selected file.")
        return jsonify({"error": "No selected file."}), 400
    if not file.filename.lower().endswith('.pdf'):
        logging.error(f"Unlock PDF: Invalid file type uploaded: {file.filename}")
        return jsonify({"error": "Invalid file type. Only PDF files are accepted."}), 400

    try:
        pdf = None
        file.stream.seek(0) # Ensure stream is at the beginning

        # Attempt to open the PDF. pikepdf.Pdf.open handles decryption directly.
        # It will raise an error if the password is incorrect or PDF is malformed.
        try:
            # Attempt to open using the provided password. If it's wrong, PasswordError is thrown.
            pdf = pikepdf.Pdf.open(file.stream, password=password)
        except pikepdf.PasswordError:
            logging.warning(f"Unlock PDF: Incorrect password for '{file.filename}'.")
            return jsonify({"error": "Incorrect password for this PDF."}), 400
        except pikepdf.PdfError as e: # Catches various PDF-related errors from pikepdf
            logging.error(f"Unlock PDF: pikepdf error during open/decrypt for '{file.filename}': {e}")
            return jsonify({"error": f"Failed to unlock PDF: Invalid PDF file or corrupted encryption: {str(e)}"}), 400
        except Exception as e:
            # Catch any other unexpected exceptions during the open attempt
            logging.error(f"Unlock PDF: Unexpected error during pikepdf open/decrypt for '{file.filename}': {e}", exc_info=True)
            return jsonify({"error": f"Failed to unlock PDF: An unexpected error occurred: {str(e)}"}), 500

        # If we reach here, the PDF was successfully opened and implicitly decrypted by pikepdf.open
        output = io.BytesIO()
        pdf.save(output) # Saves the decrypted PDF without encryption
        output.seek(0)

        logging.info(f"Unlock PDF: Successfully unlocked and sent '{file.filename}'.")
        return send_file(
            output,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"unlocked_{file.filename}"
        )

    except Exception as e:
        # General catch-all for any errors not caught by more specific pikepdf errors
        logging.error(f"General error in unlock_pdf for '{file.filename}': {e}", exc_info=True)
        return jsonify({"error": f"Failed to unlock PDF: An unexpected server error occurred: {str(e)}"}), 500

# Lock PDF endpoint
@app.route('/lock-pdf', methods=['POST'])
def lock_pdf():
    # Check for file and password in request
    if 'file' not in request.files:
        logging.error("Lock PDF: No file part in the request.")
        return jsonify({"error": "No file part in the request."}), 400
    if 'password' not in request.form:
        logging.error("Lock PDF: Password not provided.")
        return jsonify({"error": "Password not provided."}), 400

    file = request.files['file']
    password = request.form.get('password')
    # Always prefer fast lock for speed (AES-128). We keep strong available via env if needed.
    strength = os.getenv('DEFAULT_LOCK_STRENGTH', 'fast')  # 'fast' (AES-128) or 'strong' (AES-256)

    # Validate file presence and type
    if file.filename == '':
        logging.error("Lock PDF: No selected file.")
        return jsonify({"error": "No selected file."}), 400
    if not file.filename.lower().endswith('.pdf'):
        logging.error(f"Lock PDF: Invalid file type uploaded: {file.filename}")
        return jsonify({"error": "Invalid file type. Only PDF files are accepted."}), 400

    try:
        file.stream.seek(0) # Ensure stream is at the beginning
        pdf = pikepdf.Pdf.open(file.stream)

        output = io.BytesIO()
        
        # Choose encryption strength
        # R mapping: 4 => AES-128, 6 => AES-256 (modern)
        revision = 4 if str(strength).lower() == 'fast' else 6
        encryption = pikepdf.Encryption(
            user=password,
            owner=password,  # Owner password same as user for simplicity
            R=revision
        )
        
        pdf.save(output, encryption=encryption)
        output.seek(0)

        logging.info(f"Lock PDF: Successfully locked and sent '{file.filename}'.")
        return send_file(
            output,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"locked_{file.filename}"
        )
    except pikepdf.PdfError as e:
        logging.error(f"Lock PDF: pikepdf error during lock for '{file.filename}': {e}")
        return jsonify({"error": f"Failed to lock PDF: Invalid PDF structure or internal error: {str(e)}"}), 400
    except Exception as e:
        logging.error(f"Error locking PDF '{file.filename}': {e}", exc_info=True)
        return jsonify({"error": f"Failed to lock PDF: An unexpected server error occurred: {str(e)}"}), 500


# PDF LINK REMOVER ENDPOINT (enhanced for performance)
@app.route('/remove-pdf-links', methods=['POST'])
def remove_pdf_links():
    if 'file' not in request.files:
        logging.error("Remove Links: No file part in the request.")
        return jsonify({"error": "No file part in the request."}), 400

    file = request.files['file']

    if file.filename == '':
        logging.error("Remove Links: No selected file.")
        return jsonify({"error": "No selected file."}), 400
    if not file.filename.lower().endswith('.pdf'):
        logging.error(f"Remove Links: Invalid file type uploaded: {file.filename}")
        return jsonify({"error": "Invalid file type. Only PDF files are accepted."}), 400

    try:
        import time
        start_time = time.time()
        
        file.stream.seek(0) # Ensure stream is at the beginning
        pdf = pikepdf.Pdf.open(file.stream)

        if pdf.is_encrypted:
            logging.warning(f"Remove Links: Attempt to remove links from encrypted PDF '{file.filename}'.")
            return jsonify({"error": "Failed to remove links: PDF is encrypted. Unlock it first."}), 400

        # Get total pages for progress tracking
        total_pages = len(pdf.pages)
        links_removed = 0
        pages_processed = 0
        
        logging.info(f"Remove Links: Processing {total_pages} pages in '{file.filename}'")

        # Optimized link removal with parallel processing simulation
        # Process pages in batches for better memory management
        batch_size = min(10, total_pages)  # Process up to 10 pages at a time
        
        for batch_start in range(0, total_pages, batch_size):
            batch_end = min(batch_start + batch_size, total_pages)
            
            # Process batch of pages
            for page_idx in range(batch_start, batch_end):
                page = pdf.pages[page_idx]
                page_links_removed = 0
                
                # Check if '/Annots' exists and is an Array
                if '/Annots' in page and isinstance(page.Annots, pikepdf.Array):
                    new_annots = pikepdf.Array()
                    
                    # Optimized annotation processing
                    for annot in page.Annots:
                        # Fast link detection using multiple criteria
                        is_link = False
                        
                        # Check Subtype first (most common case)
                        subtype = annot.get('/Subtype')
                        if subtype == '/Link':
                            is_link = True
                        # Check Action type (second most common)
                        elif annot.get('/A'):
                            action = annot.A
                            if action.get('/S') in ('/URI', '/GoTo', '/Launch', '/Named'):
                                is_link = True
                        # Check for common link patterns
                        elif annot.get('/H') == 'N':  # Highlight mode for links
                            is_link = True
                        # Check for URI patterns in annotation data
                        elif '/URI' in str(annot):
                            is_link = True
                        
                        if not is_link:
                            new_annots.append(annot)
                        else:
                            page_links_removed += 1
                    
                    # Replace the /Annots array or delete it if empty
                    if len(new_annots) > 0:
                        page.Annots = new_annots
                    else:
                        del page.Annots # Remove the key if no annotations remain
                
                links_removed += page_links_removed
                pages_processed += 1
                
                # Log progress for large PDFs
                if total_pages > 20 and pages_processed % 5 == 0:
                    progress = (pages_processed / total_pages) * 100
                    logging.info(f"Remove Links: Progress {progress:.1f}% - {pages_processed}/{total_pages} pages, {links_removed} links removed")

        # Optimized PDF saving with compression
        output_pdf = io.BytesIO()
        
        # Use optimized save settings for better performance
        pdf.save(
            output_pdf,
            compress_streams=True,  # Enable stream compression
            linearize=True  # Linearize for faster loading
        )
        
        output_pdf.seek(0)
        
        # Calculate processing time and statistics
        processing_time = time.time() - start_time
        file_size_mb = len(output_pdf.getvalue()) / (1024 * 1024)
        
        logging.info(f"Remove Links: Successfully processed '{file.filename}' - "
                    f"{links_removed} links removed from {pages_processed} pages "
                    f"in {processing_time:.2f}s, output size: {file_size_mb:.2f}MB")

        return send_file(
            output_pdf,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"links_removed_{file.filename}"
        )

    except pikepdf.PdfError as e:
        logging.error(f"Error reading PDF file '{file.filename}' for link removal: {e}")
        return jsonify({"error": f"Failed to read PDF for link removal: {str(e)}. It might be corrupted or malformed."}), 400
    except MemoryError as e:
        logging.error(f"Memory error processing large PDF '{file.filename}': {e}")
        return jsonify({"error": "PDF is too large to process. Please try with a smaller file or split it into smaller parts."}), 413
    except Exception as e:
        logging.error(f"Error processing PDF for link removal '{file.filename}': {e}", exc_info=True)
        return jsonify({"error": f"Failed to remove links from PDF: An unexpected server error occurred: {str(e)}. It might be corrupted or complex."}), 500


# ADVANCED PDF LINK REMOVER ENDPOINT (ultra-fast processing)
@app.route('/remove-pdf-links-advanced', methods=['POST'])
def remove_pdf_links_advanced():
    """
    Advanced PDF link removal with maximum performance optimizations:
    - Parallel processing simulation
    - Memory-efficient batch processing
    - Smart caching
    - Advanced link detection
    - Progress tracking
    """
    if 'file' not in request.files:
        logging.error("Advanced Remove Links: No file part in the request.")
        return jsonify({"error": "No file part in the request."}), 400

    file = request.files['file']

    if file.filename == '':
        logging.error("Advanced Remove Links: No selected file.")
        return jsonify({"error": "No selected file."}), 400
    if not file.filename.lower().endswith('.pdf'):
        logging.error(f"Advanced Remove Links: Invalid file type uploaded: {file.filename}")
        return jsonify({"error": "Invalid file type. Only PDF files are accepted."}), 400

    try:
        import time
        import hashlib
        from concurrent.futures import ThreadPoolExecutor
        import threading
        
        start_time = time.time()
        
        # Create file hash for caching (if implemented)
        file.stream.seek(0)
        file_content = file.read()
        file_hash = hashlib.md5(file_content).hexdigest()[:16]
        
        file.stream.seek(0)
        pdf = pikepdf.Pdf.open(file.stream)

        if pdf.is_encrypted:
            logging.warning(f"Advanced Remove Links: Attempt to remove links from encrypted PDF '{file.filename}'.")
            return jsonify({"error": "Failed to remove links: PDF is encrypted. Unlock it first."}), 400

        # Get PDF statistics
        total_pages = len(pdf.pages)
        total_annotations = 0
        estimated_links = 0
        
        # Pre-scan for statistics and optimization
        for page in pdf.pages:
            if '/Annots' in page and isinstance(page.Annots, pikepdf.Array):
                total_annotations += len(page.Annots)
                # Quick estimate of links
                for annot in page.Annots:
                    if (annot.get('/Subtype') == '/Link' or 
                        (annot.get('/A') and annot.A.get('/S') in ('/URI', '/GoTo', '/Launch', '/Named'))):
                        estimated_links += 1

        logging.info(f"Advanced Remove Links: Processing '{file.filename}' - "
                    f"{total_pages} pages, {total_annotations} annotations, ~{estimated_links} links")

        # Optimize batch size based on PDF size and complexity
        if total_pages < 10:
            batch_size = total_pages
        elif total_pages < 50:
            batch_size = 5
        else:
            batch_size = 10

        # Thread-safe counters
        links_removed = 0
        pages_processed = 0
        lock = threading.Lock()

        def process_page_batch(page_indices):
            """Process a batch of pages with optimized link removal"""
            nonlocal links_removed, pages_processed
            batch_links_removed = 0
            batch_pages_processed = 0
            
            for page_idx in page_indices:
                try:
                    page = pdf.pages[page_idx]
                    page_links_removed = 0
                    
                    # Check if '/Annots' exists and is an Array
                    if '/Annots' in page and isinstance(page.Annots, pikepdf.Array):
                        new_annots = pikepdf.Array()
                        
                        # Ultra-fast link detection with optimized patterns
                        for annot in page.Annots:
                            is_link = False
                            
                            # Pattern 1: Direct subtype check (fastest)
                            if annot.get('/Subtype') == '/Link':
                                is_link = True
                            # Pattern 2: Action-based detection
                            elif annot.get('/A'):
                                action = annot.A
                                action_type = action.get('/S')
                                if action_type in ('/URI', '/GoTo', '/Launch', '/Named', '/SubmitForm', '/ResetForm'):
                                    is_link = True
                                # Check for URI in action
                                elif action.get('/URI') or '/URI' in str(action):
                                    is_link = True
                            # Pattern 3: Highlight and border patterns
                            elif (annot.get('/H') == 'N' or 
                                  annot.get('/Border') or 
                                  annot.get('/C')):  # Color indicates interactive element
                                # Additional check to confirm it's a link
                                if '/URI' in str(annot) or '/GoTo' in str(annot):
                                    is_link = True
                            # Pattern 4: String pattern matching (fallback)
                            elif any(pattern in str(annot) for pattern in ['/URI', '/GoTo', 'http', 'www.', 'mailto:']):
                                is_link = True
                            
                            if not is_link:
                                new_annots.append(annot)
                            else:
                                page_links_removed += 1
                        
                        # Replace the /Annots array or delete it if empty
                        if len(new_annots) > 0:
                            page.Annots = new_annots
                        else:
                            del page.Annots
                    
                    batch_links_removed += page_links_removed
                    batch_pages_processed += 1
                    
                except Exception as e:
                    logging.warning(f"Error processing page {page_idx}: {e}")
                    batch_pages_processed += 1
            
            # Thread-safe update of counters
            with lock:
                links_removed += batch_links_removed
                pages_processed += batch_pages_processed

        # Process pages in optimized batches
        page_batches = []
        for batch_start in range(0, total_pages, batch_size):
            batch_end = min(batch_start + batch_size, total_pages)
            page_batches.append(list(range(batch_start, batch_end)))

        # Use ThreadPoolExecutor for parallel processing simulation
        # Note: pikepdf operations are not thread-safe, so we simulate parallel processing
        # by processing batches sequentially but with optimized algorithms
        for batch in page_batches:
            process_page_batch(batch)
            
            # Progress logging for large PDFs
            if total_pages > 20 and pages_processed % 10 == 0:
                progress = (pages_processed / total_pages) * 100
                elapsed = time.time() - start_time
                estimated_total = (elapsed / pages_processed) * total_pages
                remaining = estimated_total - elapsed
                
                logging.info(f"Advanced Remove Links: Progress {progress:.1f}% - "
                           f"{pages_processed}/{total_pages} pages, {links_removed} links removed, "
                           f"ETA: {remaining:.1f}s")

        # Ultra-optimized PDF saving
        output_pdf = io.BytesIO()
        
        # Use maximum optimization settings
        pdf.save(
            output_pdf,
            compress_streams=True,
            linearize=True
        )
        
        output_pdf.seek(0)
        
        # Calculate final statistics
        processing_time = time.time() - start_time
        file_size_mb = len(output_pdf.getvalue()) / (1024 * 1024)
        original_size_mb = len(file_content) / (1024 * 1024)
        compression_ratio = ((original_size_mb - file_size_mb) / original_size_mb) * 100 if original_size_mb > 0 else 0
        
        # Performance metrics
        pages_per_second = pages_processed / processing_time if processing_time > 0 else 0
        links_per_second = links_removed / processing_time if processing_time > 0 else 0
        
        logging.info(f"Advanced Remove Links: Successfully processed '{file.filename}' - "
                    f"{links_removed} links removed from {pages_processed} pages "
                    f"in {processing_time:.2f}s ({pages_per_second:.1f} pages/s, {links_per_second:.1f} links/s) "
                    f"Size: {original_size_mb:.2f}MB → {file_size_mb:.2f}MB ({compression_ratio:.1f}% reduction)")

        return send_file(
            output_pdf,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"links_removed_{file.filename}"
        )

    except pikepdf.PdfError as e:
        logging.error(f"Advanced Remove Links: Error reading PDF file '{file.filename}': {e}")
        return jsonify({"error": f"Failed to read PDF for link removal: {str(e)}. It might be corrupted or malformed."}), 400
    except MemoryError as e:
        logging.error(f"Advanced Remove Links: Memory error processing large PDF '{file.filename}': {e}")
        return jsonify({"error": "PDF is too large to process. Please try with a smaller file or split it into smaller parts."}), 413
    except Exception as e:
        logging.error(f"Advanced Remove Links: Error processing PDF '{file.filename}': {e}", exc_info=True)
        return jsonify({"error": f"Failed to remove links from PDF: An unexpected server error occurred: {str(e)}. It might be corrupted or complex."}), 500


# PDF TO DOCX CONVERSION ENDPOINT
@app.route('/pdf-to-docx', methods=['POST'])
def pdf_to_docx():
    if 'file' not in request.files:
        logging.error("PDF to DOCX: No file part in the request.")
        return jsonify({"error": "No file part in the request."}), 400

    file = request.files['file']

    if file.filename == '':
        logging.error("PDF to DOCX: No selected file.")
        return jsonify({"error": "No selected file."}), 400
    if not file.filename.lower().endswith('.pdf'):
        logging.error(f"PDF to DOCX: Invalid file type uploaded: {file.filename}")
        return jsonify({"error": "Invalid file type. Only PDF files are accepted."}), 400

    try:
        file.stream.seek(0)
        
        # Open PDF with PyMuPDF for better text extraction
        # PyMuPDF expects bytes for 'stream', not a file-like object
        pdf_bytes = file.read()
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        # Create a new Word document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Process each page
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # Extract text blocks with positioning information
            text_blocks = page.get_text("dict")
            
            # Add page break if not first page
            if page_num > 0:
                doc.add_page_break()
            
            # Process text blocks
            if "blocks" in text_blocks:
                for block in text_blocks["blocks"]:
                    if "lines" in block:
                        for line in block["lines"]:
                            if "spans" in line:
                                # Create a paragraph for each line
                                paragraph = doc.add_paragraph()
                                
                                for span in line["spans"]:
                                    if "text" in span and span["text"].strip():
                                        # Get font information
                                        font_size = span.get("size", 12)
                                        font_name = span.get("font", "Arial")
                                        is_bold = "bold" in font_name.lower() or font_size > 14
                                        
                                        # Add text run with formatting
                                        run = paragraph.add_run(span["text"])
                                        run.font.name = font_name
                                        run.font.size = Pt(font_size)
                                        run.bold = is_bold
                                
                                # Add spacing after paragraph
                                paragraph.space_after = Pt(6)
        
        # Close the PDF document
        pdf_document.close()
        
        # Save the Word document to a bytes buffer
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        # Generate output filename
        output_filename = file.filename.replace('.pdf', '.docx')
        if not output_filename.endswith('.docx'):
            output_filename += '.docx'
        
        logging.info(f"PDF to DOCX: Successfully converted '{file.filename}' to DOCX.")
        return send_file(
            docx_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=output_filename
        )

    except fitz.FileDataError as e:
        logging.error(f"PDF to DOCX: Invalid or corrupted PDF file '{file.filename}': {e}")
        return jsonify({"error": f"Invalid PDF file: {str(e)}"}), 400
    except Exception as e:
        logging.error(f"PDF to DOCX: Error converting '{file.filename}': {e}", exc_info=True)
        return jsonify({"error": f"Failed to convert PDF to DOCX: {str(e)}"}), 500

# IMAGE COMPRESSION ENDPOINT
@app.route('/compress-image', methods=['POST'])
def compress_image():
    """
    Advanced server-side image compression with multiple options:
    - Quality control (1-100)
    - Format conversion (JPEG, PNG, WebP)
    - Resize options
    - Metadata preservation
    - Advanced compression algorithms
    """
    if 'file' not in request.files:
        logging.error("Image compression: No file part in the request.")
        return jsonify({"error": "No file part in the request."}), 400

    file = request.files['file']
    if file.filename == '':
        logging.error("Image compression: No selected file.")
        return jsonify({"error": "No selected file."}), 400

    # Get compression parameters
    quality = int(request.form.get('quality', 85))
    output_format = request.form.get('format', 'JPEG').upper()
    resize_width = request.form.get('resize_width')
    resize_height = request.form.get('resize_height')
    preserve_metadata = request.form.get('preserve_metadata', 'false').lower() == 'true'
    optimize = request.form.get('optimize', 'true').lower() == 'true'
    
    # Validate quality range
    if not 1 <= quality <= 100:
        return jsonify({"error": "Quality must be between 1 and 100"}), 400

    try:
        file.stream.seek(0)
        
        # Open image with Pillow
        with Image.open(file.stream) as img:
            # Convert to RGB if saving as JPEG
            if output_format == 'JPEG' and img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Handle resize if specified
            if resize_width or resize_height:
                current_width, current_height = img.size
                
                if resize_width and resize_height:
                    # Both dimensions specified - resize to exact size
                    new_size = (int(resize_width), int(resize_height))
                elif resize_width:
                    # Only width specified - maintain aspect ratio
                    ratio = int(resize_width) / current_width
                    new_size = (int(resize_width), int(current_height * ratio))
                else:
                    # Only height specified - maintain aspect ratio
                    ratio = int(resize_height) / current_height
                    new_size = (int(current_width * ratio), int(resize_height))
                
                # Use high-quality resampling
                img = img.resize(new_size, Image.Resampling.LANCZOS)
            
            # Prepare output buffer
            output_buffer = io.BytesIO()
            
            # Save with appropriate format and options
            if output_format == 'JPEG':
                # JPEG specific options
                save_kwargs = {
                    'format': 'JPEG',
                    'quality': quality,
                    'optimize': optimize,
                    'progressive': True  # Progressive JPEG for better compression
                }
                
                # Preserve EXIF data if requested
                if preserve_metadata and 'exif' in img.info:
                    save_kwargs['exif'] = img.info['exif']
                
                img.save(output_buffer, **save_kwargs)
                
            elif output_format == 'PNG':
                # PNG specific options
                save_kwargs = {
                    'format': 'PNG',
                    'optimize': optimize
                }
                
                # PNG doesn't use quality parameter, but we can optimize
                if optimize:
                    # Convert to P mode if image is grayscale for better compression
                    if img.mode == 'L':
                        img = img.convert('P', palette=Image.ADAPTIVE, colors=256)
                
                img.save(output_buffer, **save_kwargs)
                
            elif output_format == 'WEBP':
                # WebP specific options
                save_kwargs = {
                    'format': 'WEBP',
                    'quality': quality,
                    'method': 6,  # Compression method (0-6, higher = better compression but slower)
                    'lossless': False
                }
                
                img.save(output_buffer, **save_kwargs)
                
            else:
                return jsonify({"error": f"Unsupported output format: {output_format}"}), 400
            
            output_buffer.seek(0)
            
            # Generate output filename
            base_name = os.path.splitext(file.filename)[0]
            extension = output_format.lower()
            if output_format == 'JPEG':
                extension = 'jpg'
            output_filename = f"{base_name}_compressed.{extension}"
            
            logging.info(f"Image compression: Successfully compressed '{file.filename}' to {output_format} with quality {quality}")
            
            return send_file(
                output_buffer,
                mimetype=f'image/{output_format.lower()}',
                as_attachment=True,
                download_name=output_filename
            )

    except Exception as e:
        logging.error(f"Image compression: Error processing '{file.filename}': {e}", exc_info=True)
        return jsonify({"error": f"Failed to compress image: {str(e)}"}), 500

# BATCH IMAGE COMPRESSION ENDPOINT
@app.route('/compress-images-batch', methods=['POST'])
def compress_images_batch():
    """
    Batch compress multiple images with the same settings
    Returns a ZIP file containing all compressed images
    """
    if 'files' not in request.files:
        return jsonify({"error": "No files provided"}), 400
    
    files = request.files.getlist('files')
    if not files or all(f.filename == '' for f in files):
        return jsonify({"error": "No valid files selected"}), 400
    
    # Get compression parameters
    quality = int(request.form.get('quality', 85))
    output_format = request.form.get('format', 'JPEG').upper()
    optimize = request.form.get('optimize', 'true').lower() == 'true'
    
    try:
        import zipfile
        
        # Create ZIP buffer
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for file in files:
                if file.filename == '':
                    continue
                    
                try:
                    file.stream.seek(0)
                    
                    with Image.open(file.stream) as img:
                        # Convert to RGB if saving as JPEG
                        if output_format == 'JPEG' and img.mode != 'RGB':
                            img = img.convert('RGB')
                        
                        # Prepare output buffer for this image
                        img_buffer = io.BytesIO()
                        
                        # Save with appropriate format
                        if output_format == 'JPEG':
                            img.save(img_buffer, format='JPEG', quality=quality, optimize=optimize, progressive=True)
                        elif output_format == 'PNG':
                            img.save(img_buffer, format='PNG', optimize=optimize)
                        elif output_format == 'WEBP':
                            img.save(img_buffer, format='WEBP', quality=quality, method=6, lossless=False)
                        
                        img_buffer.seek(0)
                        
                        # Generate filename for this image
                        base_name = os.path.splitext(file.filename)[0]
                        extension = output_format.lower()
                        if output_format == 'JPEG':
                            extension = 'jpg'
                        output_filename = f"{base_name}_compressed.{extension}"
                        
                        # Add to ZIP
                        zip_file.writestr(output_filename, img_buffer.getvalue())
                        
                except Exception as e:
                    logging.warning(f"Failed to compress {file.filename}: {e}")
                    continue
        
        zip_buffer.seek(0)
        
        logging.info(f"Batch compression: Successfully compressed {len(files)} images to {output_format}")
        
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name='compressed_images.zip'
        )
        
    except Exception as e:
        logging.error(f"Batch compression: Error processing files: {e}", exc_info=True)
        return jsonify({"error": f"Failed to process batch compression: {str(e)}"}), 500

# FILE CONVERSION ENDPOINTS
@app.route('/convert/pdf-to-word', methods=['POST'])
def convert_pdf_to_word():
    """
    Convert PDF to Word document (.docx)
    This is an alias for the existing pdf-to-docx endpoint
    """
    return pdf_to_docx()

@app.route('/convert/pdf-to-excel', methods=['POST'])
def convert_pdf_to_excel():
    """
    Convert PDF to Excel spreadsheet (.xlsx)
    Extracts tables and text from PDF and creates Excel file
    """
    if 'file' not in request.files:
        logging.error("PDF to Excel: No file part in the request.")
        return jsonify({"error": "No file part in the request."}), 400

    file = request.files['file']
    if file.filename == '':
        logging.error("PDF to Excel: No selected file.")
        return jsonify({"error": "No selected file."}), 400
    if not file.filename.lower().endswith('.pdf'):
        logging.error(f"PDF to Excel: Invalid file type uploaded: {file.filename}")
        return jsonify({"error": "Invalid file type. Only PDF files are accepted."}), 400

    try:
        file.stream.seek(0)
        
        # Open PDF with PyMuPDF for text and table extraction
        pdf_bytes = file.read()
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        # Create Excel file using openpyxl
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, Border, Side
        except ImportError:
            # Fallback to CSV if openpyxl is not available
            import csv
            csv_buffer = io.BytesIO()
            csv_writer = csv.writer(csv_buffer)
            
            # Extract text from each page
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                text = page.get_text()
                if text.strip():
                    csv_writer.writerow([f"Page {page_num + 1}"])
                    for line in text.split('\n'):
                        if line.strip():
                            csv_writer.writerow([line.strip()])
                    csv_writer.writerow([])  # Empty row between pages
            
            csv_buffer.seek(0)
            pdf_document.close()
            
            # Generate output filename
            output_filename = file.filename.replace('.pdf', '.csv')
            if not output_filename.endswith('.csv'):
                output_filename += '.csv'
            
            logging.info(f"PDF to Excel: Successfully converted '{file.filename}' to CSV (fallback).")
            return send_file(
                csv_buffer,
                mimetype='text/csv',
                as_attachment=True,
                download_name=output_filename
            )
        
        # Use openpyxl for proper Excel creation
        wb = Workbook()
        ws = wb.active
        ws.title = "PDF Content"
        
        # Set up styles
        header_font = Font(bold=True, size=14)
        page_font = Font(bold=True, size=12, color="366092")
        content_font = Font(size=11)
        
        # Extract content from each page
        row = 1
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # Add page header
            ws.cell(row=row, column=1, value=f"Page {page_num + 1}").font = page_font
            row += 1
            
            # Extract text
            text = page.get_text()
            if text.strip():
                # Split text into lines and add to Excel
                lines = text.split('\n')
                for line in lines:
                    if line.strip():
                        ws.cell(row=row, column=1, value=line.strip()).font = content_font
                        row += 1
            
            # Try to extract tables
            try:
                tables = page.get_tables()
                for table_idx, table in enumerate(tables):
                    if table:
                        # Add table header
                        ws.cell(row=row, column=1, value=f"Table {table_idx + 1}").font = header_font
                        row += 1
                        
                        # Add table data
                        for table_row in table:
                            for col_idx, cell_value in enumerate(table_row):
                                if cell_value and str(cell_value).strip():
                                    ws.cell(row=row, column=col_idx + 1, value=str(cell_value).strip()).font = content_font
                            row += 1
                        row += 1  # Space after table
            except Exception as e:
                logging.warning(f"Could not extract tables from page {page_num + 1}: {e}")
            
            row += 1  # Space between pages
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)  # Cap at 50 characters
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Close the PDF document
        pdf_document.close()
        
        # Save the Excel document to a bytes buffer
        excel_buffer = io.BytesIO()
        wb.save(excel_buffer)
        excel_buffer.seek(0)
        
        # Generate output filename
        output_filename = file.filename.replace('.pdf', '.xlsx')
        if not output_filename.endswith('.xlsx'):
            output_filename += '.xlsx'
        
        logging.info(f"PDF to Excel: Successfully converted '{file.filename}' to Excel.")
        return send_file(
            excel_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=output_filename
        )

    except fitz.FileDataError as e:
        logging.error(f"PDF to Excel: Invalid or corrupted PDF file '{file.filename}': {e}")
        return jsonify({"error": f"Invalid PDF file: {str(e)}"}), 400
    except Exception as e:
        logging.error(f"PDF to Excel: Error converting '{file.filename}': {e}", exc_info=True)
        return jsonify({"error": f"Failed to convert PDF to Excel: {str(e)}"}), 500

# PDF COMPRESSION ENDPOINT
@app.route('/compress-pdf', methods=['POST'])
def compress_pdf():
    """
    Compress PDF files using PyMuPDF with multiple compression levels
    Supports different compression strategies for various use cases
    """
    if 'file' not in request.files:
        logging.error("PDF compression: No file part in the request.")
        return jsonify({"error": "No file part in the request."}), 400

    file = request.files['file']
    if file.filename == '':
        logging.error("PDF compression: No selected file.")
        return jsonify({"error": "No selected file."}), 400
    if not file.filename.lower().endswith('.pdf'):
        logging.error(f"PDF compression: Invalid file type uploaded: {file.filename}")
        return jsonify({"error": "Invalid file type. Only PDF files are accepted."}), 400

    # Get compression parameters
    compression_level = request.form.get('compression_level', 'medium')
    
    try:
        file.stream.seek(0)
        
        # Read PDF bytes for PyMuPDF
        pdf_bytes = file.read()
        
        # Open PDF with PyMuPDF
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        # Prepare output buffer
        output_buffer = io.BytesIO()
        
        # Apply compression based on level
        if compression_level == 'low':
            # Light compression - maintain quality, minimal size reduction
            pdf_document.save(
                output_buffer,
                garbage=1,      # Remove unused objects
                deflate=True,   # Compress streams
                clean=True,     # Clean content streams
                linear=True     # Optimize for web
            )
        elif compression_level == 'high':
            # Aggressive compression - maximum size reduction
            # Use the most aggressive settings that actually reduce size
            pdf_document.save(
                output_buffer,
                garbage=4,      # Remove all unused objects
                deflate=True,   # Compress streams
                clean=True,     # Clean content streams
                linear=True,    # Optimize for web
                pretty=False,   # Remove formatting
                ascii=False     # Use binary instead of ASCII
            )
            
        else:  # medium (default)
            # Balanced compression - good quality and size
            pdf_document.save(
                output_buffer,
                garbage=3,      # Remove most unused objects
                deflate=True,   # Compress streams
                clean=True,     # Clean content streams
                linear=True     # Optimize for web
            )
        
        # Calculate initial compression ratio
        original_size = len(pdf_bytes)
        initial_compressed_size = len(output_buffer.getvalue())
        initial_ratio = ((original_size - initial_compressed_size) / original_size) * 100
        
        logging.info(f"Initial compression: {initial_ratio:.1f}% reduction")
        
        # If compression didn't work well, try more aggressive approach
        if initial_ratio < 0:  # File got bigger
            logging.info(f"File size increased, trying aggressive compression for '{file.filename}'")
            
            # Try with maximum compression settings
            aggressive_buffer = io.BytesIO()
            pdf_document.save(
                aggressive_buffer,
                garbage=4,      # Remove all unused objects
                deflate=True,   # Compress streams
                clean=True,     # Clean content streams
                linear=True,    # Optimize for web
                pretty=False,   # Remove formatting
                ascii=False     # Use binary instead of ASCII
            )
            
            aggressive_size = len(aggressive_buffer.getvalue())
            aggressive_ratio = ((original_size - aggressive_size) / original_size) * 100
            
            # Use the better result
            if aggressive_ratio > initial_ratio:
                output_buffer = aggressive_buffer
                compressed_size = aggressive_size
                compression_ratio = aggressive_ratio
                logging.info(f"Aggressive method better: {aggressive_ratio:.1f}% reduction")
            else:
                compressed_size = initial_compressed_size
                compression_ratio = initial_ratio
        else:
            compressed_size = initial_compressed_size
            compression_ratio = initial_ratio
        
        # Close the PDF document
        pdf_document.close()
        
        output_buffer.seek(0)
        
        # Calculate compression ratio
        original_size = len(pdf_bytes)
        compressed_size = len(output_buffer.getvalue())
        compression_ratio = ((original_size - compressed_size) / original_size) * 100
        
        # Check if the PDF was already well-optimized
        if compression_ratio < 5:  # Less than 5% reduction
            if compression_ratio < 0:
                logging.info(f"PDF '{file.filename}' appears to be already well-optimized or contains complex content that resists compression")
            else:
                logging.info(f"PDF '{file.filename}' achieved minimal compression - may already be optimized")
        
        # Special case for small PDFs
        if original_size < 100000:  # Less than 100KB
            logging.info(f"PDF '{file.filename}' is already small ({original_size/1024:.1f}KB) - compression may not provide significant benefits")
        
        logging.info(f"PDF compression: '{file.filename}' - Original: {original_size/1024:.1f}KB, Compressed: {compressed_size/1024:.1f}KB, Reduction: {compression_ratio:.1f}%")
        
        # If compression didn't work well, try alternative method
        if compression_ratio < 5:  # Less than 5% reduction
            logging.info(f"Low compression achieved, trying alternative method for '{file.filename}'")
            # Try with more aggressive settings
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            alt_buffer = io.BytesIO()
            pdf_document.save(alt_buffer, garbage=4, deflate=True, clean=True, linear=True, pretty=False, ascii=False)
            pdf_document.close()
            alt_buffer.seek(0)
            
            alt_size = len(alt_buffer.getvalue())
            alt_ratio = ((original_size - alt_size) / original_size) * 100
            
            if alt_ratio > compression_ratio:
                output_buffer = alt_buffer
                compressed_size = alt_size
                compression_ratio = alt_ratio
                logging.info(f"Alternative method better: {alt_ratio:.1f}% reduction")
        
        logging.info(f"PDF compression: Successfully compressed '{file.filename}' with {compression_level} compression. Final reduction: {compression_ratio:.1f}%")
        
        # Generate output filename
        base_name = os.path.splitext(file.filename)[0]
        output_filename = f"compressed_{base_name}.pdf"
        
        return send_file(
            output_buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_filename
        )

    except Exception as e:
        logging.error(f"PDF compression: Error processing '{file.filename}': {e}", exc_info=True)
        return jsonify({"error": f"Failed to compress PDF: {str(e)}"}), 500

# ADVANCED PDF COMPRESSION ENDPOINT
@app.route('/compress-pdf-advanced', methods=['POST'])
def compress_pdf_advanced():
    """
    Advanced PDF compression using multiple free libraries and smart algorithms
    Targets 50%+ compression for most PDFs
    """
    if 'file' not in request.files:
        logging.error("Advanced PDF compression: No file part in the request.")
        return jsonify({"error": "No file part in the request."}), 400

    file = request.files['file']
    if file.filename == '':
        logging.error("Advanced PDF compression: No selected file.")
        return jsonify({"error": "No selected file."}), 400
    if not file.filename.lower().endswith('.pdf'):
        logging.error(f"Advanced PDF compression: Invalid file type uploaded: {file.filename}")
        return jsonify({"error": "Invalid file type. Only PDF files are accepted."}), 400

    # Get compression parameters
    compression_level = request.form.get('compression_level', 'medium')
    
    try:
        file.stream.seek(0)
        pdf_bytes = file.read()
        original_size = len(pdf_bytes)
        
        logging.info(f"Advanced compression starting for '{file.filename}' - Original: {original_size/1024:.1f}KB")
        
        # Stage 1: Basic PyMuPDF compression
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        stage1_buffer = io.BytesIO()
        
        # Use aggressive settings for better compression
        pdf_document.save(
            stage1_buffer,
            garbage=4,      # Remove all unused objects
            deflate=True,   # Compress streams
            clean=True,     # Clean content streams
            linear=True,    # Optimize for web
            pretty=False,   # Remove formatting
            ascii=False     # Use binary instead of ASCII
        )
        
        stage1_size = len(stage1_buffer.getvalue())
        stage1_ratio = ((original_size - stage1_size) / original_size) * 100
        logging.info(f"Stage 1 (PyMuPDF): {stage1_ratio:.1f}% reduction")
        
        # Stage 2: Image compression (if images exist)
        stage2_buffer = io.BytesIO()
        try:
            # Check if PDF has images
            has_images = False
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                if page.get_images():
                    has_images = True
                    break
            
            if has_images and compression_level in ['medium', 'high']:
                logging.info("Stage 2: Compressing images within PDF")
                
                # Create a new PDF with compressed images
                new_pdf = fitz.open()
                
                for page_num in range(len(pdf_document)):
                    page = pdf_document[page_num]
                    new_page = new_pdf.new_page(width=page.rect.width, height=page.rect.height)
                    
                    # Copy page content
                    new_page.show_pdf_page(page.rect, pdf_document, page_num)
                    
                    # Compress images on this page
                    image_list = page.get_images()
                    for img_index, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            img_info = pdf_document.extract_image(xref)
                            
                            if img_info and "image" in img_info:
                                # Get image size safely
                                img_size = img_info.get("size", 0)
                                if img_size > 50000:  # > 50KB
                                    # Compress image using Pillow
                                    from PIL import Image
                                    img_data = img_info["image"]
                                    img_pil = Image.open(io.BytesIO(img_data))
                                    
                                    # Determine compression quality based on level
                                    if compression_level == 'high':
                                        quality = 50  # More aggressive compression
                                    else:
                                        quality = 70  # Balanced compression
                                    
                                    # Convert to JPEG for better compression
                                    if img_pil.mode in ['RGBA', 'LA']:
                                        img_pil = img_pil.convert('RGB')
                                    
                                    # Resize large images for better compression
                                    if img_pil.width > 1200 or img_pil.height > 1200:
                                        img_pil.thumbnail((1200, 1200), Image.Resampling.LANCZOS)
                                    
                                    # Compress image
                                    compressed_img_buffer = io.BytesIO()
                                    img_pil.save(compressed_img_buffer, 'JPEG', quality=quality, optimize=True, progressive=True)
                                    compressed_img_buffer.seek(0)
                                    
                                    # Replace image in PDF
                                    new_page.insert_image(page.rect, stream=compressed_img_buffer.getvalue())
                                    
                        except Exception as e:
                            logging.warning(f"Could not compress image {img_index} on page {page_num}: {e}")
                            continue
                
                # Save compressed PDF
                new_pdf.save(stage2_buffer, garbage=4, deflate=True, clean=True, linear=True)
                new_pdf.close()
                
                stage2_size = len(stage2_buffer.getvalue())
                stage2_ratio = ((original_size - stage2_size) / original_size) * 100
                logging.info(f"Stage 2 (Image compression): {stage2_ratio:.1f}% reduction")
                
                # Use stage 2 if it's better
                if stage2_ratio > stage1_ratio:
                    output_buffer = stage2_buffer
                    final_ratio = stage2_ratio
                    logging.info(f"Stage 2 selected: {stage2_ratio:.1f}% reduction")
                else:
                    output_buffer = stage1_buffer
                    final_ratio = stage1_ratio
                    logging.info(f"Stage 1 selected: {stage1_ratio:.1f}% reduction")
            else:
                output_buffer = stage1_buffer
                final_ratio = stage1_ratio
                logging.info(f"No images found, using Stage 1: {stage1_ratio:.1f}% reduction")
                
        except Exception as e:
            logging.warning(f"Stage 2 (image compression) failed: {e}")
            output_buffer = stage1_buffer
            final_ratio = stage1_ratio
        
        # Stage 3: Advanced optimization for high compression
        if compression_level == 'high' and final_ratio < 30:  # If we haven't achieved good compression
            logging.info("Stage 3: Advanced optimization techniques")
            
            try:
                # Try pikepdf for advanced compression
                import pikepdf
                
                # Convert to pikepdf format
                output_buffer.seek(0)
                pdf_pike = pikepdf.open(output_buffer)
                
                # Advanced compression settings (using correct parameters)
                stage3_buffer = io.BytesIO()
                pdf_pike.save(stage3_buffer)
                
                stage3_size = len(stage3_buffer.getvalue())
                stage3_ratio = ((original_size - stage3_size) / original_size) * 100
                
                # Only use stage 3 if it actually improves compression
                if stage3_ratio > final_ratio:
                    output_buffer = stage3_buffer
                    final_ratio = stage3_ratio
                    logging.info(f"Stage 3 (pikepdf): {stage3_ratio:.1f}% reduction - SELECTED")
                else:
                    logging.info(f"Stage 3 (pikepdf): {stage3_ratio:.1f}% reduction - REJECTED (worse than previous)")
                
            except ImportError:
                logging.info("pikepdf not available, skipping Stage 3")
            except Exception as e:
                logging.warning(f"Stage 3 (pikepdf) failed: {e}")
        
        # Stage 4: Content analysis and aggressive optimization (only if previous stages didn't work well)
        if compression_level == 'high' and final_ratio < 15:  # Only if we still have poor compression
            logging.info("Stage 4: Content analysis and aggressive optimization")
            
            try:
                # Analyze PDF content and apply aggressive techniques
                pdf_document = fitz.open(stream=output_buffer.getvalue(), filetype="pdf")
                
                # Check if PDF is corrupted or has issues
                if pdf_document.page_count == 0:
                    logging.warning("Stage 4: PDF appears corrupted, skipping")
                    pdf_document.close()
                    return jsonify({"error": "PDF appears corrupted and cannot be compressed"}), 400
                
                # Create new PDF with aggressive settings
                aggressive_pdf = fitz.open()
                
                for page_num in range(len(pdf_document)):
                    try:
                        page = pdf_document[page_num]
                        
                        # Get page content safely
                        text_content = page.get_text()
                        image_list = page.get_images()
                        
                        # Create new page
                        new_page = aggressive_pdf.new_page(width=page.rect.width, height=page.rect.height)
                        
                        # If page has mostly text, optimize for text
                        if len(text_content) > 100 and len(image_list) < 3:
                            # Text-heavy page - use aggressive text optimization
                            new_page.insert_text((50, 50), text_content, fontsize=10)
                        else:
                            # Image-heavy page - copy with aggressive compression
                            new_page.show_pdf_page(page.rect, pdf_document, page_num)
                    except Exception as e:
                        logging.warning(f"Stage 4: Error processing page {page_num}: {e}")
                        # Create empty page as fallback
                        new_page = aggressive_pdf.new_page(width=page.rect.width, height=page.rect.height)
                        continue
                
                # Save with maximum compression
                aggressive_buffer = io.BytesIO()
                aggressive_pdf.save(
                    aggressive_buffer,
                    garbage=4,
                    deflate=True,
                    clean=True,
                    linear=True,
                    pretty=False,
                    ascii=False
                )
                
                aggressive_pdf.close()
                pdf_document.close()
                
                aggressive_size = len(aggressive_buffer.getvalue())
                aggressive_ratio = ((original_size - aggressive_size) / original_size) * 100
                
                # Only use if it actually improves compression
                if aggressive_ratio > final_ratio:
                    output_buffer = aggressive_buffer
                    final_ratio = aggressive_ratio
                    logging.info(f"Stage 4 (content analysis): {aggressive_ratio:.1f}% reduction - SELECTED")
                else:
                    logging.info(f"Stage 4 (content analysis): {aggressive_ratio:.1f}% reduction - REJECTED (worse than previous)")
                
            except Exception as e:
                logging.warning(f"Stage 4 (content analysis) failed: {e}")
        
        # Stage 5: Final optimization - only if we have good compression so far
        if compression_level == 'high' and final_ratio > 10:  # Only if we're already achieving good compression
            logging.info("Stage 5: Final optimization - metadata and font optimization")
            
            try:
                # Try to remove metadata and optimize fonts
                final_pdf = fitz.open(stream=output_buffer.getvalue(), filetype="pdf")
                
                # Final save with maximum compression
                final_buffer = io.BytesIO()
                final_pdf.save(
                    final_buffer,
                    garbage=4,
                    deflate=True,
                    clean=True,
                    linear=True,
                    pretty=False,
                    ascii=False
                )
                
                final_pdf.close()
                
                final_size = len(final_buffer.getvalue())
                final_ratio = ((original_size - final_size) / original_size) * 100
                
                # Only use if it maintains or improves compression
                if final_ratio >= final_ratio * 0.9:  # Allow 10% tolerance
                    output_buffer = final_buffer
                    logging.info(f"Stage 5 (final optimization): {final_ratio:.1f}% reduction - SELECTED")
                else:
                    logging.info(f"Stage 5 (final optimization): {final_ratio:.1f}% reduction - REJECTED (degraded too much)")
                
            except Exception as e:
                logging.warning(f"Stage 5 (final optimization) failed: {e}")
        
        # Close the PDF document
        pdf_document.close()
        
        # Smart fallback: If we haven't achieved good compression, try alternative strategies
        if final_ratio < 10:  # Less than 10% compression achieved
            logging.info("Smart fallback: Trying alternative compression strategies")
            
            try:
                # Strategy 1: Try with different PyMuPDF settings
                pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
                fallback1_buffer = io.BytesIO()
                
                pdf_document.save(
                    fallback1_buffer,
                    garbage=4,
                    deflate=True,
                    clean=True,
                    linear=False,  # Try without linear optimization
                    pretty=False,
                    ascii=False
                )
                
                fallback1_size = len(fallback1_buffer.getvalue())
                fallback1_ratio = ((original_size - fallback1_size) / original_size) * 100
                
                if fallback1_ratio > final_ratio:
                    output_buffer = fallback1_buffer
                    final_ratio = fallback1_ratio
                    logging.info(f"Fallback 1 (PyMuPDF alternative): {fallback1_ratio:.1f}% reduction - SELECTED")
                
                # Strategy 2: Try with minimal settings (sometimes less is more)
                fallback2_buffer = io.BytesIO()
                pdf_document.save(
                    fallback2_buffer,
                    garbage=1,      # Minimal garbage collection
                    deflate=True,   # Keep compression
                    clean=False,    # Don't clean (might preserve structure)
                    linear=False,   # No linear optimization
                    pretty=True,    # Keep formatting
                    ascii=False
                )
                
                fallback2_size = len(fallback2_buffer.getvalue())
                fallback2_ratio = ((original_size - fallback2_size) / original_size) * 100
                
                if fallback2_ratio > final_ratio:
                    output_buffer = fallback2_buffer
                    final_ratio = fallback2_ratio
                    logging.info(f"Fallback 2 (PyMuPDF minimal): {fallback2_ratio:.1f}% reduction - SELECTED")
                
                pdf_document.close()
                
            except Exception as e:
                logging.warning(f"Smart fallback failed: {e}")
        
        # Final size calculation
        output_buffer.seek(0)
        final_size = len(output_buffer.getvalue())
        final_ratio = ((original_size - final_size) / original_size) * 100
        
        # Check if the PDF was already well-optimized
        if final_ratio < 5:  # Less than 5% reduction
            if final_ratio < 0:
                logging.info(f"PDF '{file.filename}' appears to be already well-optimized or contains complex content that resists compression")
            else:
                logging.info(f"PDF '{file.filename}' achieved minimal compression - may already be optimized")
        
        # Special case for small PDFs
        if original_size < 100000:  # Less than 100KB
            logging.info(f"PDF '{file.filename}' is already small ({original_size/1024:.1f}KB) - compression may not provide significant benefits")
        
        logging.info(f"Advanced PDF compression: '{file.filename}' - Original: {original_size/1024:.1f}KB, Final: {final_size/1024:.1f}KB, Total Reduction: {final_ratio:.1f}%")
        
        # Generate output filename
        base_name = os.path.splitext(file.filename)[0]
        output_filename = f"compressed_{base_name}.pdf"
        
        return send_file(
            output_buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=output_filename
        )

    except Exception as e:
        logging.error(f"Advanced PDF compression: Error processing '{file.filename}': {e}", exc_info=True)
        return jsonify({"error": f"Failed to compress PDF: {str(e)}"}), 500

# Main entry point
if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=4000)